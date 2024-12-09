import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import *
from tkinter.ttk import Combobox
from tkinter import ttk
from tkinter import messagebox
from tkinter.ttk import Radiobutton


root = Tk()
root.geometry('650x450')
root.title("Калькулятор транспортного налога")
root['bg'] = '#c7dcd0'
root.resizable(width=False, height=False)

def create_window():
    try:
        year = year_choose.get()
        months = int(number_of_months_choose.get())
        vehicle_type = type_of_vehicle_choose.get()
        try:
            power = float(txt_power.get())
        except ValueError:
            messagebox.showerror("Ошибка", "Некорректный формат мощности!")
            return

        allowed_ranges = {
            "Автомобили легковые": (0, 2000),
            "Мотоциклы и мотороллеры": (0, 100),
            "Автобусы": (0, 500),
            "Автомобили грузовые": (0, 750),
            "Другие самоходные транспортные средства": (0, 1000),
            "Машины и механизмы на пневматическом и гусеничном ходу (с каждой лошадиной силы)": (0, 1000),
            "Снегоходы, мотосани": (0, 500),
            "Катера, моторные лодки и другие водные транспортные средства": (0, 500),
            "Яхты и другие парусно-моторные суда": (0, 500),
            "Несамоходные (буксируемые) суда": (0, 500),
            "Самолеты, вертолеты и иные воздушные суда, имеющие двигатели (с каждой лошадиной силы)": (0, 1000),
            "Самолеты с реактивными двигателями (с каждого килограмма силы тяги)": (0, 1000),
            "Другие водные и воздушные транспортные средства без двигателей (с единицы транспортного средства)": (0, 1000),
        }
        min_power, max_power = allowed_ranges.get(vehicle_type, (0, float('inf')))

        if not (min_power <= power <= max_power):
            messagebox.showerror("Ошибка", f"Мощность вне допустимого диапазона для {vehicle_type} (от {min_power} до {max_power} л.с.)")
            return


        expensive_car = expensive_car_var.get() if expensive_car_checkbutton.winfo_exists() else False

        if not all([year, months, vehicle_type, power]):
            messagebox.showerror("Ошибка", "Заполните все поля!")
            return

        tax = calculate_tax(vehicle_type, power, months, expensive_car)

        window = Tk()
        window.title("Результат расчета")
        window.geometry("650x350")
        window.resizable(width=False, height=False)

        lbl1 = Label(window, text="Расчет носит справочный характер и не может быть основанием для совершения юридически значимых действий.", font=("Times new roman", 8))
        lbl1.grid(column=0, row=0)
        lbl2 = Label(window, text="Расчёт транспортного налога", font=("Times new roman", 18))
        lbl2.grid(column=0, row=1)
        lbl3 = Label(window, text=f"Сумма транспортного налога: {tax:.2f} руб.", font=("Times new roman", 18))
        lbl3.grid(column=0, row=2, sticky=W)
        lbl4 = Label(window, text=f"Вид ТС: {vehicle_type}", font=("Times new roman", 14))
        lbl4.grid(column=0, row=3, sticky=W)
        lbl5 = Label(window, text=f"Мощность ТС: {power} л.с.", font=("Times new roman", 14))
        lbl5.grid(column=0, row=4, sticky=W)
        lbl6 = Label(window, text=f"Год расчета и количество месяцев владения: {year} год, {months} месяцев", font=("Times new roman", 14))
        lbl6.grid(column=0, row=5, sticky=W)
        lbl7 = Label(window, text=f"Ставка: {get_tax_rate(vehicle_type, power)} руб./л.с.", font=("Times new roman", 14))
        lbl7.grid(column=0, row=6, sticky=W)
        lbl8 = Label(window, text="Порядок расчёта", font=("Times new roman", 18))
        lbl8.grid(column=0, row=7, sticky=W)
        k_coefficient = 3 if expensive_car else 1
        lbl9 = Label(window, text=f"{get_tax_rate(vehicle_type, power)} * {power} * {months}/12 * {k_coefficient} = {tax:.2f}", font=("Times new roman", 14))
        lbl9.grid(column=0, row=8)
        btn = Button(window, text="Сформировать отчет", command=lambda: create_doc(tax, vehicle_type, power, year, months, expensive_car))
        btn.grid(column=0, row=9)
        window.mainloop()

    except ValueError:
        messagebox.showerror("Ошибка", "Некорректный формат данных!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

def create_doc(tax, vehicle_type, power, year, months, expensive_car):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times new roman'
    style.font.size = Pt(14)

    part1 = doc.add_paragraph('Расчет носит справочный характер и не может быть основанием для совершения юридически значимых действий.')
    run = part1.runs[0]
    run.font.size = Pt(8)
    part1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part2 = doc.add_paragraph('Расчёт транспортного налога')
    run1 = part2.runs[0]
    run1.font.size = Pt(18)
    part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    part3 = doc.add_paragraph(f'Сумма транспортного налога: {tax:.2f} руб.')
    run2 = part3.runs[0]
    run2.font.size = Pt(18)
    part4 = doc.add_paragraph(f"Вид ТС: {vehicle_type}")
    part5 = doc.add_paragraph(f"Мощность ТС: {power} л.с.")
    part6 = doc.add_paragraph(f"Год расчета и количество месяцев владения: {year} год, {months} месяцев")
    part7 = doc.add_paragraph(f"Ставка: {get_tax_rate(vehicle_type, power)} руб./л.с.")
    part8 = doc.add_paragraph("Порядок расчёта")
    run3 = part8.runs[0]
    run3.font.size = Pt(18)
    k_coefficient = 3 if expensive_car else 1
    part9 = doc.add_paragraph(f"{get_tax_rate(vehicle_type, power)} * {power} * {months}/12 * {k_coefficient} = {tax:.2f}")
    doc.save('Отчет.docx')


def calculate_tax(vehicle_type, power, months, expensive_car):
    k = 1
    if expensive_car:
        k = 3
    r = get_tax_rate(vehicle_type, power)
    if r is None:
        return "Ошибка: Неверный тип транспортного средства или мощность."
    return r * power * (months / 12) * k


def get_tax_rate(vehicle_type, power):
    if vehicle_type == "Автомобили легковые":
        if 0 <= power <= 100: return 25
        elif 100 < power <= 150: return 35
        elif 150 < power <= 200: return 50
        elif 200 < power <= 250: return 75
        elif 250 < power <= 2000: return 150
        else: return None
    elif vehicle_type == "Мотоциклы и мотороллеры":
        if 0 <= power <= 20: return 10
        elif 20 < power <= 35: return 20
        elif 35 < power <= 100: return 50
        else: return None
    elif vehicle_type == "Автобусы":
        if 0 <= power <= 200: return 50
        elif 200 < power <= 500: return 100
        else: return None
    elif vehicle_type == "Автомобили грузовые":
        if 0 <= power <= 100: return 25
        elif 100 < power <= 150: return 40
        elif 150 < power <= 200: return 50
        elif 200 < power <= 250: return 65
        elif 250 < power <= 750: return 85
        else: return None
    elif vehicle_type == "Другие самоходные транспортные средства":
        if 0 <= power <= 1000: return 25
        else: return None
    elif vehicle_type == "Машины и механизмы на пневматическом и гусеничном ходу (с каждой лошадиной силы)":
        if 0 <= power <= 1000: return 25
        else: return None
    elif vehicle_type == "Снегоходы, мотосани":
        if 0 <= power <= 50: return 25
        elif 50 < power <= 500: return 50
        else: return None
    elif vehicle_type == "Катера, моторные лодки и другие водные транспортные средства":
        if 0 <= power <= 100: return 50
        elif 100 < power <= 500: return 100
        else: return None
    elif vehicle_type == "Яхты и другие парусно-моторные суда":
        if 0 <= power <= 100: return 125
        elif 100 < power <= 500: return 250
        else: return None
    elif vehicle_type == "Несамоходные (буксируемые) суда":
        if 0 <= power <= 500: return 100
        else: return None
    elif vehicle_type == "Самолеты, вертолеты и иные воздушные суда, имеющие двигатели (с каждой лошадиной силы)":
        if 0 <= power <= 1000: return 125
        else: return None
    elif vehicle_type == "Самолеты с реактивными двигателями (с каждого килограмма силы тяги)":
        if 0 <= power <= 1000: return 100
        else: return None
    elif vehicle_type == "Другие водные и воздушные транспортные средства без двигателей (с единицы транспортного средства)":
        if 0 <= power <= 1000: return 1000
        else: return None
    else: return None


year = Label(root, text="Выберите год", bg='#c7dcd0')
year.grid(column=0, row=0, sticky=W)
year_choose = Combobox(root, state="readonly")
year_choose['values'] = (2021, 2022, 2023, 2024)
year_choose.current(3)
year_choose.grid(column=1, row=0, sticky=W)

number_of_months = Label(root, text="Выберите количество\n месяцев владения ТС", bg='#c7dcd0')
number_of_months.grid(column=0, row=1, sticky=W)
number_of_months_choose = Combobox(root, state="readonly")
number_of_months_choose['values'] = (1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12)
number_of_months_choose.current(11)
number_of_months_choose.grid(column=1, row=1, sticky=W)

type_of_vehicle = Label(root, text="Выберите вид\n транспортного средства", bg='#c7dcd0')
type_of_vehicle.grid(column=0, row=2, sticky=W)
type_of_vehicle_choose = Combobox(root, state="readonly",)
type_of_vehicle_choose['values'] = (
    "           ",
    "Автомобили легковые",
    "Мотоциклы и мотороллеры",
    "Автобусы",
    "Автомобили грузовые",
    "Другие самоходные транспортные средства",
    "Машины и механизмы на пневматическом и гусеничном ходу (с каждой лошадиной силы)",
    "Снегоходы, мотосани",
    "Катера, моторные лодки и другие водные транспортные средства",
    "Яхты и другие парусно-моторные суда",
    "Гидроциклы",
    "Несамоходные (буксируемые) суда",
    "Самолеты, вертолеты и иные воздушные суда, имеющие двигатели (с каждой лошадиной силы)",
    "Самолеты с реактивными двигателями (с каждого килограмма силы тяги)",
    "Другие водные и воздушные транспортные средства без двигателей (с единицы транспортного средства)",
)
type_of_vehicle_choose.current(0)
type_of_vehicle_choose.grid(column=1, row=2, sticky=W)

engine_power = Label(root, text="Введите мощность\n двигателя (л.с.)", bg='#c7dcd0')
engine_power.grid(column=0, row=3, sticky=W)
txt_power = Entry(root, width=23)
txt_power.grid(column=1, row=3, sticky=W)

expensive_car_var = BooleanVar()
expensive_car_checkbutton = Checkbutton(root, text='Стоимость автомобиля больше 10 млн.руб.?', variable=expensive_car_var)

def show_hide_expensive_car_checkbox(event):
    if type_of_vehicle_choose.get() == "Автомобили легковые":
        expensive_car_checkbutton.grid(column=2, row=2, sticky=W)
    else:
        expensive_car_checkbutton.grid_remove()
        expensive_car_var.set(False)

type_of_vehicle_choose.bind("<<ComboboxSelected>>", show_hide_expensive_car_checkbox)

btn = Button(root, text="Рассчитать налог", command=create_window)
btn.grid(column=1, row=4, sticky=W)

root.mainloop()
